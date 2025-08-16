import os
import hashlib
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, BinaryIO
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import logging
import mimetypes
from io import BytesIO

# Optional dependencies with fallbacks
try:
    import magic
except ImportError:
    magic = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

class DocumentFormat(Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOCX = "docx" 
    TXT = "txt"
    MARKDOWN = "md"
    JPEG = "jpg"
    PNG = "png"

class ValidationStatus(Enum):
    """Document validation status"""
    VALID = "valid"
    INVALID = "invalid"
    SUSPICIOUS = "suspicious"
    UNSUPPORTED = "unsupported"

@dataclass
class IngestionConfig:
    """Configuration for document ingestion"""
    max_file_size_mb: int = 100
    supported_formats: List[str] = field(default_factory=lambda: ['pdf', 'docx', 'txt', 'md', 'jpg', 'png'])
    enable_content_validation: bool = True
    enable_malware_check: bool = True
    extract_preview: bool = True
    preview_length: int = 1000
    enable_format_conversion: bool = False

@dataclass
class DocumentMetadata:
    """Comprehensive document metadata"""
    file_name: str
    file_path: str
    file_size: int
    file_hash: str
    mime_type: str
    detected_format: DocumentFormat
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    language: Optional[str] = None

@dataclass
class ValidationResult:
    """Document validation result"""
    status: ValidationStatus
    is_valid: bool
    confidence_score: float
    issues: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    security_flags: List[str] = field(default_factory=list)

@dataclass
class IngestedDocument:
    """Complete ingested document with metadata"""
    content: bytes
    metadata: DocumentMetadata
    validation_result: ValidationResult
    preview_text: Optional[str] = None
    processing_timestamp: datetime = field(default_factory=datetime.utcnow)
    ingestion_source: str = "unknown"

@dataclass
class NormalizedDocument:
    """Normalized document ready for processing"""
    original_document: IngestedDocument
    normalized_content: bytes
    content_type: str
    processing_hints: Dict[str, Any] = field(default_factory=dict)
    quality_score: float = 1.0

class DocumentIngestionManager:
    """Production-ready document ingestion with comprehensive format support"""
    
    # MIME type mappings
    MIME_TYPE_MAPPINGS = {
        'application/pdf': DocumentFormat.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentFormat.DOCX,
        'text/plain': DocumentFormat.TXT,
        'text/markdown': DocumentFormat.MARKDOWN,
        'image/jpeg': DocumentFormat.JPEG,
        'image/png': DocumentFormat.PNG
    }
    
    # File extension mappings
    EXTENSION_MAPPINGS = {
        '.pdf': DocumentFormat.PDF,
        '.docx': DocumentFormat.DOCX,
        '.txt': DocumentFormat.TXT,
        '.md': DocumentFormat.MARKDOWN,
        '.jpg': DocumentFormat.JPEG,
        '.jpeg': DocumentFormat.JPEG,
        '.png': DocumentFormat.PNG
    }
    
    def __init__(self, config: IngestionConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Initialize file type detection
        self.magic = None
        if magic:
            try:
                self.magic = magic.Magic(mime=True)
            except Exception as e:
                self.logger.warning(f"libmagic not available, falling back to mimetypes: {e}")
        else:
            self.logger.info("python-magic not installed, using mimetypes for format detection")
    
    def ingest_document(self, file_content: bytes, file_metadata: Dict[str, Any]) -> IngestedDocument:
        """Ingest single document with comprehensive processing"""
        
        try:
            # Extract basic file information
            file_path = file_metadata.get('file_path', 'unknown')
            file_name = Path(file_path).name
            
            self.logger.info(f"Ingesting document: {file_name}")
            
            # Validate file size
            if len(file_content) > self.config.max_file_size_mb * 1024 * 1024:
                raise ValueError(f"File size {len(file_content)} exceeds limit of {self.config.max_file_size_mb}MB")
            
            # Detect document format
            detected_format = self._detect_format(file_content, file_path)
            
            # Validate format support
            if detected_format.value not in self.config.supported_formats:
                raise ValueError(f"Unsupported format: {detected_format.value}")
            
            # Extract comprehensive metadata
            doc_metadata = self.extract_metadata(file_content, file_path)
            doc_metadata.detected_format = detected_format
            
            # Validate document content
            validation_result = self.validate_document(file_content, detected_format.value)
            
            # Extract preview text
            preview_text = None
            if self.config.extract_preview:
                preview_text = self._extract_preview(file_content, detected_format)
            
            # Create ingested document
            ingested_doc = IngestedDocument(
                content=file_content,
                metadata=doc_metadata,
                validation_result=validation_result,
                preview_text=preview_text,
                ingestion_source=file_metadata.get('source', 'gitlab')
            )
            
            self.logger.info(f"Successfully ingested: {file_name} ({detected_format.value}, {len(file_content)} bytes)")
            return ingested_doc
            
        except Exception as e:
            self.logger.error(f"Document ingestion failed for {file_path}: {e}")
            raise
    
    def _detect_format(self, file_content: bytes, file_path: str) -> DocumentFormat:
        """Detect document format using multiple methods"""
        
        # Method 1: MIME type detection
        mime_type = None
        if self.magic:
            try:
                mime_type = self.magic.from_buffer(file_content)
            except Exception as e:
                self.logger.warning(f"Magic MIME detection failed: {e}")
        
        if not mime_type:
            # Fallback to mimetypes module
            mime_type, _ = mimetypes.guess_type(file_path)
        
        # Map MIME type to format
        if mime_type in self.MIME_TYPE_MAPPINGS:
            return self.MIME_TYPE_MAPPINGS[mime_type]
        
        # Method 2: File extension
        file_ext = Path(file_path).suffix.lower()
        if file_ext in self.EXTENSION_MAPPINGS:
            return self.EXTENSION_MAPPINGS[file_ext]
        
        # Method 3: Content-based detection
        return self._detect_format_by_content(file_content)
    
    def _detect_format_by_content(self, file_content: bytes) -> DocumentFormat:
        """Detect format by analyzing file content"""
        
        # Check for PDF signature
        if file_content.startswith(b'%PDF'):
            return DocumentFormat.PDF
        
        # Check for ZIP signature (DOCX is a ZIP file)
        if file_content.startswith(b'PK\x03\x04'):
            # Could be DOCX, check for Office Open XML structure
            if b'word/' in file_content[:1024] or b'[Content_Types].xml' in file_content[:1024]:
                return DocumentFormat.DOCX
        
        # Check for JPEG signature
        if file_content.startswith(b'\xff\xd8\xff'):
            return DocumentFormat.JPEG
        
        # Check for PNG signature
        if file_content.startswith(b'\x89PNG\r\n\x1a\n'):
            return DocumentFormat.PNG
        
        # Check if content is text-based
        try:
            text_content = file_content.decode('utf-8')
            # Simple heuristic for markdown
            if any(marker in text_content[:500] for marker in ['# ', '## ', '* ', '- ', '```']):
                return DocumentFormat.MARKDOWN
            else:
                return DocumentFormat.TXT
        except UnicodeDecodeError:
            pass
        
        # Default fallback
        raise ValueError("Unable to detect document format")
    
    def validate_document(self, file_content: bytes, expected_format: str) -> ValidationResult:
        """Comprehensive document validation"""
        
        issues = []
        warnings = []
        security_flags = []
        confidence_score = 1.0
        
        try:
            format_enum = DocumentFormat(expected_format)
            
            # Format-specific validation
            if format_enum == DocumentFormat.PDF:
                issues.extend(self._validate_pdf(file_content))
            elif format_enum == DocumentFormat.DOCX:
                issues.extend(self._validate_docx(file_content))
            elif format_enum in [DocumentFormat.JPEG, DocumentFormat.PNG]:
                issues.extend(self._validate_image(file_content))
            elif format_enum in [DocumentFormat.TXT, DocumentFormat.MARKDOWN]:
                issues.extend(self._validate_text(file_content))
            
            # General security checks
            if self.config.enable_malware_check:
                security_flags.extend(self._basic_security_check(file_content))
            
            # Calculate confidence score
            if issues:
                confidence_score = max(0.0, 1.0 - (len(issues) * 0.2))
            
            # Determine overall status
            if security_flags:
                status = ValidationStatus.SUSPICIOUS
                is_valid = False
            elif issues:
                status = ValidationStatus.INVALID
                is_valid = False
            else:
                status = ValidationStatus.VALID
                is_valid = True
            
            return ValidationResult(
                status=status,
                is_valid=is_valid,
                confidence_score=confidence_score,
                issues=issues,
                warnings=warnings,
                security_flags=security_flags
            )
            
        except Exception as e:
            self.logger.error(f"Document validation failed: {e}")
            return ValidationResult(
                status=ValidationStatus.INVALID,
                is_valid=False,
                confidence_score=0.0,
                issues=[f"Validation error: {e}"]
            )
    
    def _validate_pdf(self, file_content: bytes) -> List[str]:
        """Validate PDF document"""
        issues = []
        
        if not fitz:
            issues.append("PyMuPDF not available for PDF validation")
            return issues
        
        try:
            # Open PDF with PyMuPDF
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            # Check if PDF is valid
            if doc.is_pdf:
                # Check for password protection
                if doc.needs_pass:
                    issues.append("PDF is password protected")
                
                # Check page count
                if doc.page_count == 0:
                    issues.append("PDF has no pages")
                elif doc.page_count > 1000:
                    issues.append("PDF has unusually high page count")
                
                # Check for corruption indicators
                try:
                    for page_num in range(min(3, doc.page_count)):  # Check first 3 pages
                        page = doc[page_num]
                        text = page.get_text()
                        if not text.strip() and not page.get_images():
                            issues.append(f"Page {page_num + 1} appears to be empty")
                except Exception as e:
                    issues.append(f"Error reading PDF content: {e}")
            else:
                issues.append("File is not a valid PDF")
            
            doc.close()
            
        except Exception as e:
            issues.append(f"PDF validation error: {e}")
        
        return issues
    
    def _validate_docx(self, file_content: bytes) -> List[str]:
        """Validate DOCX document"""
        issues = []
        
        if not DocxDocument:
            issues.append("python-docx not available for DOCX validation")
            return issues
        
        try:
            doc = DocxDocument(BytesIO(file_content))
            
            # Check if document has content
            paragraph_count = len(doc.paragraphs)
            if paragraph_count == 0:
                issues.append("DOCX document has no paragraphs")
            
            # Check for corrupted content
            try:
                text_content = '\n'.join([p.text for p in doc.paragraphs[:10]])  # Check first 10 paragraphs
                if not text_content.strip():
                    issues.append("DOCX document appears to have no readable text")
            except Exception as e:
                issues.append(f"Error reading DOCX content: {e}")
                
        except Exception as e:
            issues.append(f"DOCX validation error: {e}")
        
        return issues
    
    def _validate_image(self, file_content: bytes) -> List[str]:
        """Validate image document"""
        issues = []
        
        if not Image:
            issues.append("Pillow not available for image validation")
            return issues
        
        try:
            image = Image.open(BytesIO(file_content))
            
            # Check image properties
            if image.size[0] < 100 or image.size[1] < 100:
                issues.append("Image resolution is very low")
            
            if image.size[0] > 10000 or image.size[1] > 10000:
                issues.append("Image resolution is unusually high")
            
            # Verify image can be processed
            image.verify()
            
        except Exception as e:
            issues.append(f"Image validation error: {e}")
        
        return issues
    
    def _validate_text(self, file_content: bytes) -> List[str]:
        """Validate text document"""
        issues = []
        
        try:
            # Try to decode as UTF-8
            text_content = file_content.decode('utf-8')
            
            # Check content length
            if len(text_content.strip()) == 0:
                issues.append("Text document is empty")
            elif len(text_content) < 10:
                issues.append("Text document is very short")
            
            # Check for unusual characters
            non_printable_count = sum(1 for c in text_content if ord(c) < 32 and c not in '\n\r\t')
            if non_printable_count > len(text_content) * 0.1:
                issues.append("Text contains many non-printable characters")
                
        except UnicodeDecodeError as e:
            issues.append(f"Text encoding error: {e}")
        except Exception as e:
            issues.append(f"Text validation error: {e}")
        
        return issues
    
    def _basic_security_check(self, file_content: bytes) -> List[str]:
        """Basic security checks for malicious content"""
        security_flags = []
        
        # Check for suspicious patterns (very basic)
        suspicious_patterns = [
            b'<script',
            b'javascript:',
            b'vbscript:',
            b'onclick=',
            b'onerror=',
            b'eval(',
            b'exec(',
        ]
        
        content_lower = file_content[:10000].lower()  # Check first 10KB
        for pattern in suspicious_patterns:
            if pattern in content_lower:
                security_flags.append(f"Suspicious pattern detected: {pattern.decode('utf-8', errors='ignore')}")
        
        # Check file size anomalies
        if len(file_content) > 50 * 1024 * 1024:  # 50MB
            security_flags.append("File size is unusually large")
        
        return security_flags
    
    def extract_metadata(self, file_content: bytes, file_path: str) -> DocumentMetadata:
        """Extract comprehensive document metadata"""
        
        file_name = Path(file_path).name
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Basic metadata
        metadata = DocumentMetadata(
            file_name=file_name,
            file_path=file_path,
            file_size=len(file_content),
            file_hash=file_hash,
            mime_type="",
            detected_format=DocumentFormat.TXT  # Will be updated
        )
        
        # Detect MIME type
        if self.magic:
            try:
                metadata.mime_type = self.magic.from_buffer(file_content)
            except Exception:
                metadata.mime_type, _ = mimetypes.guess_type(file_path) or ("", "")
        else:
            metadata.mime_type, _ = mimetypes.guess_type(file_path) or ("", "")
        
        # Format-specific metadata extraction
        try:
            format_enum = self._detect_format(file_content, file_path)
            metadata.detected_format = format_enum
            
            if format_enum == DocumentFormat.PDF:
                self._extract_pdf_metadata(file_content, metadata)
            elif format_enum == DocumentFormat.DOCX:
                self._extract_docx_metadata(file_content, metadata)
            elif format_enum in [DocumentFormat.TXT, DocumentFormat.MARKDOWN]:
                self._extract_text_metadata(file_content, metadata)
        
        except Exception as e:
            self.logger.warning(f"Metadata extraction failed for {file_path}: {e}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_content: bytes, metadata: DocumentMetadata):
        """Extract PDF-specific metadata"""
        if not fitz:
            return
        
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            # Basic properties
            metadata.page_count = doc.page_count
            
            # Document metadata
            doc_metadata = doc.metadata
            if doc_metadata:
                metadata.title = doc_metadata.get('title', '')
                metadata.author = doc_metadata.get('author', '')
                
                # Parse creation date
                creation_date = doc_metadata.get('creationDate', '')
                if creation_date:
                    try:
                        # PDF dates are in format: D:YYYYMMDDHHmmSSOHH'mm
                        if creation_date.startswith('D:'):
                            date_str = creation_date[2:16]  # YYYYMMDDHHMMSS
                            metadata.created_date = datetime.strptime(date_str, '%Y%m%d%H%M%S')
                    except Exception:
                        pass
            
            # Extract text for word count
            try:
                all_text = ""
                for page_num in range(min(5, doc.page_count)):  # First 5 pages for estimation
                    page_text = doc[page_num].get_text()
                    all_text += page_text
                
                if all_text:
                    metadata.word_count = len(all_text.split())
                    metadata.character_count = len(all_text)
            except Exception:
                pass
            
            doc.close()
            
        except Exception as e:
            self.logger.warning(f"PDF metadata extraction failed: {e}")
    
    def _extract_docx_metadata(self, file_content: bytes, metadata: DocumentMetadata):
        """Extract DOCX-specific metadata"""
        if not DocxDocument:
            return
        
        try:
            doc = DocxDocument(BytesIO(file_content))
            
            # Core properties
            core_props = doc.core_properties
            metadata.title = core_props.title or ''
            metadata.author = core_props.author or ''
            metadata.created_date = core_props.created
            metadata.modified_date = core_props.modified
            
            # Text statistics
            all_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            metadata.word_count = len(all_text.split())
            metadata.character_count = len(all_text)
            
        except Exception as e:
            self.logger.warning(f"DOCX metadata extraction failed: {e}")
    
    def _extract_text_metadata(self, file_content: bytes, metadata: DocumentMetadata):
        """Extract text-specific metadata"""
        try:
            text_content = file_content.decode('utf-8')
            
            metadata.word_count = len(text_content.split())
            metadata.character_count = len(text_content)
            
            # Try to detect language (simple heuristic)
            if any(char in text_content for char in 'çãõáéíóúâêîôûà'):
                metadata.language = 'pt-BR'
            elif any(word in text_content.lower() for word in ['the', 'and', 'or', 'but']):
                metadata.language = 'en-US'
            
        except Exception as e:
            self.logger.warning(f"Text metadata extraction failed: {e}")
    
    def _extract_preview(self, file_content: bytes, doc_format: DocumentFormat) -> Optional[str]:
        """Extract preview text from document"""
        
        try:
            if doc_format == DocumentFormat.PDF:
                return self._extract_pdf_preview(file_content)
            elif doc_format == DocumentFormat.DOCX:
                return self._extract_docx_preview(file_content)
            elif doc_format in [DocumentFormat.TXT, DocumentFormat.MARKDOWN]:
                return self._extract_text_preview(file_content)
            else:
                return f"[{doc_format.value.upper()} file - no text preview available]"
        
        except Exception as e:
            self.logger.warning(f"Preview extraction failed: {e}")
            return None
    
    def _extract_pdf_preview(self, file_content: bytes) -> str:
        """Extract preview text from PDF"""
        if not fitz:
            return "[PDF preview not available - PyMuPDF not installed]"
        
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            preview_text = ""
            for page_num in range(min(3, doc.page_count)):  # First 3 pages
                page_text = doc[page_num].get_text()
                preview_text += page_text
                
                if len(preview_text) >= self.config.preview_length:
                    break
            
            doc.close()
            return preview_text[:self.config.preview_length]
            
        except Exception as e:
            return f"[PDF preview extraction failed: {e}]"
    
    def _extract_docx_preview(self, file_content: bytes) -> str:
        """Extract preview text from DOCX"""
        if not DocxDocument:
            return "[DOCX preview not available - python-docx not installed]"
        
        try:
            doc = DocxDocument(BytesIO(file_content))
            
            preview_text = ""
            for paragraph in doc.paragraphs:
                preview_text += paragraph.text + "\n"
                
                if len(preview_text) >= self.config.preview_length:
                    break
            
            return preview_text[:self.config.preview_length]
            
        except Exception as e:
            return f"[DOCX preview extraction failed: {e}]"
    
    def _extract_text_preview(self, file_content: bytes) -> str:
        """Extract preview text from text files"""
        try:
            text_content = file_content.decode('utf-8')
            return text_content[:self.config.preview_length]
            
        except Exception as e:
            return f"[Text preview extraction failed: {e}]"
    
    def normalize_content(self, document: IngestedDocument) -> NormalizedDocument:
        """Normalize document content for downstream processing"""
        
        try:
            # Basic normalization - prepare content for Docling
            normalized_content = document.content
            content_type = document.metadata.detected_format.value
            processing_hints = {}
            quality_score = document.validation_result.confidence_score
            
            # Format-specific processing hints
            if document.metadata.detected_format == DocumentFormat.PDF:
                processing_hints['requires_ocr'] = self._pdf_requires_ocr(document.content)
                processing_hints['page_count'] = document.metadata.page_count
            elif document.metadata.detected_format == DocumentFormat.DOCX:
                processing_hints['has_tables'] = True  # Assume tables might exist
                processing_hints['structured_content'] = True
            elif document.metadata.detected_format in [DocumentFormat.JPEG, DocumentFormat.PNG]:
                processing_hints['requires_ocr'] = True
                processing_hints['image_processing'] = True
            else:
                processing_hints['text_only'] = True
            
            # Quality assessment
            if document.validation_result.issues:
                quality_score = max(0.5, quality_score - 0.1 * len(document.validation_result.issues))
            
            return NormalizedDocument(
                original_document=document,
                normalized_content=normalized_content,
                content_type=content_type,
                processing_hints=processing_hints,
                quality_score=quality_score
            )
            
        except Exception as e:
            self.logger.error(f"Document normalization failed: {e}")
            raise
    
    def _pdf_requires_ocr(self, file_content: bytes) -> bool:
        """Determine if PDF requires OCR processing"""
        if not fitz:
            return True  # Default to requiring OCR if we can't analyze
        
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            # Check first few pages for text content
            for page_num in range(min(3, doc.page_count)):
                page = doc[page_num]
                text = page.get_text().strip()
                
                # If we find meaningful text, OCR is not required
                if len(text) > 50:  # Threshold for meaningful text
                    doc.close()
                    return False
            
            doc.close()
            return True  # No meaningful text found, requires OCR
            
        except Exception:
            return True  # Default to requiring OCR on error
    
    def batch_ingest(self, documents: List[Dict[str, Any]]) -> List[IngestedDocument]:
        """Batch process multiple documents"""
        
        ingested_docs = []
        errors = []
        
        self.logger.info(f"Starting batch ingestion of {len(documents)} documents")
        
        for i, doc_info in enumerate(documents):
            try:
                file_content = doc_info['content']
                file_metadata = doc_info['metadata']
                
                ingested_doc = self.ingest_document(file_content, file_metadata)
                ingested_docs.append(ingested_doc)
                
                self.logger.debug(f"Document {i+1}/{len(documents)} processed successfully")
                
            except Exception as e:
                error_msg = f"Batch ingestion failed for document {i+1}: {e}"
                self.logger.error(error_msg)
                errors.append(error_msg)
                # Continue with other documents
        
        success_rate = len(ingested_docs) / len(documents) * 100 if documents else 0
        self.logger.info(f"Batch ingestion completed: {len(ingested_docs)}/{len(documents)} documents processed ({success_rate:.1f}% success rate)")
        
        if errors:
            self.logger.warning(f"Batch ingestion had {len(errors)} errors")
        
        return ingested_docs
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats"""
        return self.config.supported_formats.copy()
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics and capabilities"""
        return {
            'supported_formats': self.get_supported_formats(),
            'max_file_size_mb': self.config.max_file_size_mb,
            'capabilities': {
                'pdf_processing': fitz is not None,
                'docx_processing': DocxDocument is not None,
                'image_processing': Image is not None,
                'mime_detection': self.magic is not None,
                'content_validation': self.config.enable_content_validation,
                'malware_check': self.config.enable_malware_check,
                'preview_extraction': self.config.extract_preview
            }
        }

def create_document_ingestion_manager(config_dict: Dict[str, Any]) -> DocumentIngestionManager:
    """Factory function for document ingestion manager creation"""
    config = IngestionConfig(
        max_file_size_mb=config_dict.get('max_file_size_mb', 100),
        supported_formats=config_dict.get('supported_formats', ['pdf', 'docx', 'txt', 'md', 'jpg', 'png']),
        enable_content_validation=config_dict.get('enable_content_validation', True),
        enable_malware_check=config_dict.get('enable_malware_check', True),
        extract_preview=config_dict.get('extract_preview', True),
        preview_length=config_dict.get('preview_length', 1000),
        enable_format_conversion=config_dict.get('enable_format_conversion', False)
    )
    return DocumentIngestionManager(config)

# Context manager for document ingestion
class DocumentIngestionContext:
    """Context manager for document ingestion with automatic cleanup"""
    
    def __init__(self, config_dict: Dict[str, Any]):
        self.config_dict = config_dict
        self.manager = None
    
    def __enter__(self) -> DocumentIngestionManager:
        self.manager = create_document_ingestion_manager(self.config_dict)
        return self.manager
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        # Cleanup if needed
        pass

# Document ingestion error classes
class DocumentIngestionError(Exception):
    """Base exception for document ingestion errors"""
    pass

class DocumentValidationError(DocumentIngestionError):
    """Document validation failures"""
    pass

class UnsupportedFormatError(DocumentIngestionError):
    """Unsupported document format errors"""
    pass

class DocumentSizeError(DocumentIngestionError):
    """Document size limit exceeded errors"""
    pass

class SecurityCheckError(DocumentIngestionError):
    """Security validation failures"""
    pass