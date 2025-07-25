"""
DOCX Document Generator

This module implements the DOCX document generator using python-docx,
providing Microsoft Word document generation with full formatting support.
"""

import io
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
import asyncio

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_ORIENT
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")

from ..base import BaseDocumentGenerator, DocumentFormat, GenerationError
from ..models import DocumentContent, DocumentTemplate, GenerationOptions, ContentType, ContentElement
from .docx.styles import DocxStyleManager
from .docx.elements import DocxElementHandler
from .docx.images import DocxImageHandler
from .docx.tables import DocxTableHandler

logger = logging.getLogger(__name__)


class DocxGenerator(BaseDocumentGenerator):
    """
    DOCX document generator using python-docx.
    
    This generator creates Microsoft Word documents from structured content
    with support for rich formatting, images, tables, and corporate templates.
    """
    
    def __init__(self):
        """Initialize the DOCX generator."""
        super().__init__(DocumentFormat.DOCX)
        
        # Initialize component handlers
        self.style_manager = DocxStyleManager()
        self.element_handler = DocxElementHandler()
        self.image_handler = DocxImageHandler()
        self.table_handler = DocxTableHandler()
        
        # Document creation settings
        self._current_document: Optional[Document] = None
        self._template_document: Optional[Document] = None
    
    @property
    def supported_features(self) -> List[str]:
        """Return list of features supported by this generator."""
        return [
            "text", "paragraphs", "headings", "lists", "tables", 
            "images", "styles", "templates", "metadata", "headers_footers",
            "page_breaks", "bookmarks", "hyperlinks", "fonts", "colors"
        ]
    
    @property
    def file_extension(self) -> str:
        """Return the file extension for DOCX format."""
        return "docx"
    
    @property
    def mime_type(self) -> str:
        """Return the MIME type for DOCX format."""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    async def generate_document(
        self,
        content: DocumentContent,
        template: Optional[DocumentTemplate] = None,
        options: Optional[GenerationOptions] = None
    ) -> bytes:
        """
        Generate a DOCX document from content and template.
        
        Args:
            content: Structured document content
            template: Optional template to apply
            options: Generation options and settings
            
        Returns:
            Generated DOCX document as bytes
            
        Raises:
            GenerationError: If document generation fails
        """
        try:
            logger.info("Starting DOCX document generation")
            
            # Validate inputs
            validation_errors = await self.validate_content(content)
            if validation_errors:
                raise GenerationError(f"Content validation failed: {', '.join(validation_errors)}")
            
            # Set generation options
            if options:
                self.set_generation_options(options)
            
            # Create or load document from template
            if template and template.template_file_path:
                self._current_document = Document(template.template_file_path)
                logger.info(f"Loaded template: {template.template_file_path}")
            else:
                self._current_document = Document()
                logger.info("Created new document")
            
            # Apply template styles if provided
            if template:
                await self._apply_template(template)
            
            # Set document properties and metadata
            await self._set_document_metadata(content.metadata)
            
            # Configure page setup
            await self._configure_page_setup(content)
            
            # Generate document content
            await self._generate_content(content)
            
            # Save document to bytes
            doc_buffer = io.BytesIO()
            self._current_document.save(doc_buffer)
            doc_bytes = doc_buffer.getvalue()
            
            logger.info(f"DOCX document generated successfully ({len(doc_bytes)} bytes)")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            raise GenerationError(f"Failed to generate DOCX document: {str(e)}") from e
        finally:
            self._current_document = None
            self._template_document = None
    
    async def generate_preview(
        self,
        content: DocumentContent,
        template: Optional[DocumentTemplate] = None,
        options: Optional[GenerationOptions] = None
    ) -> Dict[str, Any]:
        """
        Generate a preview representation of the DOCX document.
        
        Args:
            content: Structured document content
            template: Optional template to apply
            options: Generation options and settings
            
        Returns:
            Preview data with structure information
        """
        try:
            # Create document structure preview without full rendering
            preview_data = {
                "format": "docx",
                "page_count": self._estimate_page_count(content),
                "sections": [],
                "images": [],
                "tables": [],
                "metadata": {
                    "title": content.metadata.title,
                    "author": content.metadata.author,
                    "word_count": self._estimate_word_count(content),
                    "created": content.created_timestamp.isoformat()
                },
                "template": template.name if template else "Default",
                "features_used": self._analyze_content_features(content)
            }
            
            # Analyze sections for preview
            for i, section in enumerate(content.sections):
                section_preview = {
                    "index": i,
                    "title": section.title,
                    "level": section.level,
                    "element_count": len(section.elements),
                    "elements": []
                }
                
                # Preview first few elements
                for j, element in enumerate(section.elements[:5]):  # Preview first 5 elements
                    element_preview = {
                        "type": element.type.value,
                        "content_preview": element.content[:100] + "..." if len(element.content) > 100 else element.content
                    }
                    section_preview["elements"].append(element_preview)
                
                preview_data["sections"].append(section_preview)
            
            return preview_data
            
        except Exception as e:
            logger.error(f"Preview generation failed: {str(e)}")
            return {
                "format": "docx",
                "error": str(e),
                "page_count": 0,
                "sections": []
            }
    
    async def validate_content(self, content: DocumentContent) -> List[str]:
        """
        Validate content compatibility with DOCX generator.
        
        Args:
            content: Document content to validate
            
        Returns:
            List of validation errors (empty if valid)
        """
        errors = []
        
        # Basic content validation
        if not content.sections:
            errors.append("Document must contain at least one section")
        
        # Check for unsupported elements
        for i, section in enumerate(content.sections):
            for j, element in enumerate(section.elements):
                if element.type == ContentType.IMAGE:
                    if not element.image_data:
                        errors.append(f"Section {i}, Element {j}: Image element missing image_data")
                    # Note: Missing or invalid image sources are handled gracefully with placeholders
                
                elif element.type == ContentType.TABLE:
                    if not element.table_data or not element.table_data.rows:
                        errors.append(f"Section {i}, Element {j}: Table missing data or rows")
        
        return errors
    
    async def estimate_generation_time(self, content: DocumentContent) -> float:
        """
        Estimate DOCX generation time in seconds.
        
        Args:
            content: Document content to analyze
            
        Returns:
            Estimated generation time in seconds
        """
        # Base time for document creation
        base_time = 0.5
        
        # Time per element (rough estimates)
        element_count = sum(len(section.elements) for section in content.sections)
        element_time = element_count * 0.02  # 20ms per element
        
        # Additional time for images
        image_count = sum(
            1 for section in content.sections
            for element in section.elements
            if element.type == ContentType.IMAGE
        )
        image_time = image_count * 0.5  # 500ms per image
        
        # Additional time for tables
        table_count = sum(
            1 for section in content.sections
            for element in section.elements
            if element.type == ContentType.TABLE
        )
        table_time = table_count * 0.2  # 200ms per table
        
        total_time = base_time + element_time + image_time + table_time
        return max(total_time, 0.1)  # Minimum 100ms
    
    async def _apply_template(self, template: DocumentTemplate) -> None:
        """Apply template styles and configuration to document."""
        logger.info(f"Applying template: {template.name}")
        
        # Apply styles through style manager
        await self.style_manager.apply_template_styles(self._current_document, template)
        
        # Configure page layout
        if template.layout:
            await self._apply_template_layout(template.layout)
    
    async def _apply_template_layout(self, layout) -> None:
        """Apply template layout configuration."""
        if not self._current_document:
            return
            
        section = self._current_document.sections[0]
        
        # Set page size and orientation
        if layout.page_size == "A4":
            if layout.page_orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)
            else:
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
        elif layout.page_size == "Letter":
            if layout.page_orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11)
                section.page_height = Inches(8.5)
            else:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
        
        # Set margins
        if layout.margins:
            section.top_margin = Cm(layout.margins.get("top", 2.54))
            section.bottom_margin = Cm(layout.margins.get("bottom", 2.54))
            section.left_margin = Cm(layout.margins.get("left", 2.54))
            section.right_margin = Cm(layout.margins.get("right", 2.54))
    
    async def _set_document_metadata(self, metadata) -> None:
        """Set document metadata and properties."""
        if not self._current_document or not metadata:
            return
            
        core_props = self._current_document.core_properties
        
        if metadata.title:
            core_props.title = metadata.title
        if metadata.author:
            core_props.author = metadata.author
        if metadata.subject:
            core_props.subject = metadata.subject
        if metadata.description:
            core_props.comments = metadata.description
        if metadata.keywords:
            core_props.keywords = "; ".join(metadata.keywords)
        if metadata.category:
            core_props.category = metadata.category
        
        # Set dates
        if metadata.created_date:
            core_props.created = metadata.created_date
        if metadata.modified_date:
            core_props.modified = metadata.modified_date
        
        core_props.language = metadata.language
        core_props.version = metadata.version
    
    async def _configure_page_setup(self, content: DocumentContent) -> None:
        """Configure page setup based on content settings."""
        if not self._current_document:
            return
            
        section = self._current_document.sections[0]
        
        # Set page orientation
        if content.page_orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        
        # Set margins if specified
        if content.margins:
            section.top_margin = Cm(content.margins.get("top", 2.54))
            section.bottom_margin = Cm(content.margins.get("bottom", 2.54))
            section.left_margin = Cm(content.margins.get("left", 2.54))
            section.right_margin = Cm(content.margins.get("right", 2.54))
    
    async def _generate_content(self, content: DocumentContent) -> None:
        """Generate the main document content."""
        if not self._current_document:
            return
            
        # Set document title if provided
        if content.metadata.title:
            title_paragraph = self._current_document.add_heading(content.metadata.title, level=0)
        
        # Process each section
        for section in content.sections:
            await self._generate_section(section)
    
    async def _generate_section(self, section) -> None:
        """Generate content for a document section."""
        # Add section title if provided
        if section.title:
            heading = self._current_document.add_heading(section.title, level=section.level)
        
        # Add page break before if specified
        if section.page_break_before:
            self._current_document.add_page_break()
        
        # Process section elements
        for element in section.elements:
            await self._generate_element(element)
        
        # Process subsections recursively
        for subsection in section.subsections:
            await self._generate_section(subsection)
        
        # Add page break after if specified
        if section.page_break_after:
            self._current_document.add_page_break()
    
    async def _generate_element(self, element: ContentElement) -> None:
        """Generate content for a single element."""
        if element.type == ContentType.HEADING:
            await self.element_handler.add_heading(self._current_document, element)
        elif element.type == ContentType.PARAGRAPH:
            await self.element_handler.add_paragraph(self._current_document, element)
        elif element.type == ContentType.LIST:
            await self.element_handler.add_list(self._current_document, element)
        elif element.type == ContentType.TABLE:
            await self.table_handler.add_table(self._current_document, element)
        elif element.type == ContentType.IMAGE:
            await self.image_handler.add_image(self._current_document, element)
        elif element.type == ContentType.CODE:
            await self.element_handler.add_code_block(self._current_document, element)
        elif element.type == ContentType.QUOTE:
            await self.element_handler.add_quote(self._current_document, element)
        elif element.type == ContentType.DIVIDER:
            await self.element_handler.add_divider(self._current_document, element)
        else:
            # Default to paragraph for unknown types
            await self.element_handler.add_paragraph(self._current_document, element)
    
    def _estimate_page_count(self, content: DocumentContent) -> int:
        """Estimate page count for preview."""
        # Rough estimation based on content length
        total_chars = sum(
            len(element.content) 
            for section in content.sections 
            for element in section.elements
        )
        
        # Assume ~2000 characters per page (rough estimate)
        estimated_pages = max(1, total_chars // 2000)
        
        # Add pages for images and tables
        image_count = sum(
            1 for section in content.sections
            for element in section.elements
            if element.type == ContentType.IMAGE
        )
        table_count = sum(
            1 for section in content.sections
            for element in section.elements
            if element.type == ContentType.TABLE
        )
        
        # Images and tables typically take more space
        estimated_pages += (image_count // 2) + (table_count // 3)
        
        return estimated_pages
    
    def _estimate_word_count(self, content: DocumentContent) -> int:
        """Estimate word count for preview."""
        total_words = 0
        for section in content.sections:
            for element in section.elements:
                if element.content:
                    # Simple word count (split by whitespace)
                    total_words += len(element.content.split())
        return total_words
    
    def _analyze_content_features(self, content: DocumentContent) -> List[str]:
        """Analyze what features are used in the content."""
        features = set()
        
        for section in content.sections:
            for element in section.elements:
                features.add(element.type.value)
        
        return list(features)


# Test functionality
async def test_docx_generator():
    """Test DOCX generator functionality."""
    from ..models import DocumentContent, DocumentMetadata, DocumentSection, create_heading, create_paragraph
    
    # Create test content
    metadata = DocumentMetadata(
        title="Test DOCX Document",
        author="Test Author",
        subject="Testing DOCX Generation"
    )
    
    section = DocumentSection(
        title="Test Section",
        elements=[
            create_heading("Test Heading", level=1),
            create_paragraph("This is a test paragraph for DOCX generation.")
        ]
    )
    
    content = DocumentContent(
        metadata=metadata,
        sections=[section]
    )
    
    # Test generator
    generator = DocxGenerator()
    
    # Test validation
    errors = await generator.validate_content(content)
    if errors:
        print(f"❌ Validation failed: {errors}")
        return False
    
    # Test preview generation
    preview = await generator.generate_preview(content)
    print(f"✅ Preview generated: {preview['page_count']} pages, {preview['metadata']['word_count']} words")
    
    # Test document generation
    try:
        doc_bytes = await generator.generate_document(content)
        print(f"✅ Document generated: {len(doc_bytes)} bytes")
        return True
    except Exception as e:
        print(f"❌ Document generation failed: {e}")
        return False


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == "test":
        asyncio.run(test_docx_generator())
    else:
        print("Usage: python -m src.generation.formats.docx_generator test")